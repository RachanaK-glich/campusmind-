import os
import re
import logging
from typing import List, Dict, Any

logger = logging.getLogger("campusmind.extractor")


class DocumentPage:
    def __init__(self, page_number: int, text: str):
        self.page_number = page_number
        self.text = text


class DocumentExtractor:
    @staticmethod
    def extract_from_file(file_path: str) -> List[DocumentPage]:
        """Extract text from PDF, DOCX, or TXT file into a list of DocumentPage objects."""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")

        ext = os.path.splitext(file_path)[1].lower()
        if ext == ".pdf":
            return DocumentExtractor._extract_pdf(file_path)
        elif ext in [".docx", ".doc"]:
            return DocumentExtractor._extract_docx(file_path)
        elif ext in [".txt", ".md", ".csv", ".json"]:
            return DocumentExtractor._extract_text(file_path)
        else:
            # Fallback to text reading
            return DocumentExtractor._extract_text(file_path)

    @staticmethod
    def _clean_text(text: str) -> str:
        """Clean and normalize extracted text."""
        if not text:
            return ""
        # Normalize whitespace while preserving paragraph breaks
        text = re.sub(r"\r\n|\r", "\n", text)
        text = re.sub(r"[ \t]+", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()

    @staticmethod
    def _extract_pdf(file_path: str) -> List[DocumentPage]:
        """Extract text page-by-page from PDF."""
        pages: List[DocumentPage] = []
        
        # Try PyPDF first
        try:
            from pypdf import PdfReader
            reader = PdfReader(file_path)
            for idx, page in enumerate(reader.pages):
                extracted = page.extract_text() or ""
                cleaned = DocumentExtractor._clean_text(extracted)
                if cleaned:
                    pages.append(DocumentPage(page_number=idx + 1, text=cleaned))
            if pages:
                return pages
        except Exception as e:
            logger.warning(f"PyPDF extraction error: {e}. Trying fallback.")

        # Fallback to PyMuPDF if available
        try:
            import fitz
            doc = fitz.open(file_path)
            for idx, page in enumerate(doc):
                extracted = page.get_text() or ""
                cleaned = DocumentExtractor._clean_text(extracted)
                if cleaned:
                    pages.append(DocumentPage(page_number=idx + 1, text=cleaned))
            if pages:
                return pages
        except Exception as e:
            logger.warning(f"PyMuPDF extraction note: {e}")

        # If empty, return at least one page with placeholder or error
        if not pages:
            pages.append(DocumentPage(page_number=1, text="[No readable text found in PDF document]"))
        return pages

    @staticmethod
    def _extract_docx(file_path: str) -> List[DocumentPage]:
        """Extract text from DOCX document."""
        try:
            import docx
            doc = docx.Document(file_path)
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text.strip())
            
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        full_text.append(row_text)

            combined = "\n\n".join(full_text)
            cleaned = DocumentExtractor._clean_text(combined)
            
            # Estimate pages roughly (~1800 chars per page)
            if len(cleaned) <= 2000:
                return [DocumentPage(page_number=1, text=cleaned)]
            
            chunks = []
            lines = cleaned.split("\n\n")
            current_page = 1
            current_text = []
            current_len = 0
            
            for line in lines:
                current_text.append(line)
                current_len += len(line)
                if current_len >= 1800:
                    chunks.append(DocumentPage(page_number=current_page, text="\n\n".join(current_text)))
                    current_page += 1
                    current_text = []
                    current_len = 0
            
            if current_text:
                chunks.append(DocumentPage(page_number=current_page, text="\n\n".join(current_text)))
            return chunks
        except Exception as e:
            logger.error(f"DOCX extraction error: {e}")
            return [DocumentPage(page_number=1, text=f"Error extracting DOCX: {str(e)}")]

    @staticmethod
    def _extract_text(file_path: str) -> List[DocumentPage]:
        """Extract text from plain text or markdown file."""
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()
            cleaned = DocumentExtractor._clean_text(content)
            
            # Check for page markers like --- Page 2 --- or split every 2000 chars
            page_splits = re.split(r"(?:---|===)\s*(?:Page|PAGE)\s*(\d+)\s*(?:---|===)", cleaned)
            if len(page_splits) > 1:
                pages = []
                # First chunk before any marker
                if page_splits[0].strip():
                    pages.append(DocumentPage(page_number=1, text=page_splits[0].strip()))
                for i in range(1, len(page_splits), 2):
                    page_num = int(page_splits[i])
                    text_content = page_splits[i+1].strip() if i+1 < len(page_splits) else ""
                    if text_content:
                        pages.append(DocumentPage(page_number=page_num, text=text_content))
                return pages

            # If no page markers, split into logical 2000 char pages
            if len(cleaned) <= 2200:
                return [DocumentPage(page_number=1, text=cleaned)]
            
            pages = []
            paragraphs = cleaned.split("\n\n")
            curr_text = []
            curr_len = 0
            page_no = 1
            
            for p in paragraphs:
                curr_text.append(p)
                curr_len += len(p)
                if curr_len >= 1800:
                    pages.append(DocumentPage(page_number=page_no, text="\n\n".join(curr_text)))
                    page_no += 1
                    curr_text = []
                    curr_len = 0
            
            if curr_text:
                pages.append(DocumentPage(page_number=page_no, text="\n\n".join(curr_text)))
            return pages
        except Exception as e:
            logger.error(f"Text file read error: {e}")
            return [DocumentPage(page_number=1, text=f"Error reading file: {str(e)}")]
